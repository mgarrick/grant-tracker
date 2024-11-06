import os
import sys

import pandas as pd
import docx.oxml.ns as ns
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re
import html
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openai import AzureOpenAI
from datetime import datetime

from pandas import NaT
import argparse

API_KEY_FILE = 'apikey.txt'
FEW_SHOT_EXAMPLES_CSV = 'extra/few_shot_examples.csv'


# 1. read csv file
def read_csv_file(file_path):
    try:
        # Read the CSV file into a DataFrame
        df = pd.read_csv(file_path)
        # Keep only the specified fields (columns)
        fields = ["Title", "Funder", "Deadline", "Amount", "Eligibility", "Career Stage", "Abstract", "More Information"]
        filtered_df = df[fields]
        # remove html tag in each cell
        filtered_df = filtered_df.map(remove_html_tag)
        # Return the DataFrame
        return filtered_df
    except FileNotFoundError:
        print("File not found!")
        return None
    except KeyError:
        print("Some specified fields do not exist in the CSV file!")
        return None
    except Exception as e:
        print(f"An error occurred: {e}")
        return None

def remove_html_tag(value):
    if not isinstance(value, str):
        value = str(value)

    # Unescape HTML entities like &nbsp;, &amp;, etc.
    value = html.unescape(value)

    # Remove HTML tags
    value = re.sub(r"<[^>]*>", "", value)

    return value


# 2. convert csv file to word file and format word file
def format_word_file(llm, data_frame, head_title):

    # Create a new Word document for the formatted content
    formatted_doc = Document()
    
    # Add head_title as the first line, centered and bold
    if head_title:
        title = formatted_doc.add_paragraph(head_title)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.bold = True
        title_run.font.size = Pt(14)

    data_frame['ClosestFutureDate'] = data_frame['Deadline'].apply(extract_closest_future_date)
    # Sort the Deadline
    data_frame = data_frame.sort_values(by='ClosestFutureDate')
    
    # Iterate over each row in the DataFrame, skipping the first row
    for _, row in data_frame.iterrows():
        p = formatted_doc.add_paragraph()

        # title
        title_text = f"{row['Funder']} | {row['Title']}"
        if row['More Information']:
            hyperlink_run = add_hyperlink(p, row['More Information'], title_text)
            p.add_run("\n")
        else:
            title_run = p.add_run(title_text)
            title_run.bold = True
            p.add_run("\n")

        # Deadline
        if row['Deadline']:
            deadline_txt = row['Deadline']
            bold_run = p.add_run(f"Due Date: ")
            bold_run.bold = True
            closest_future_date = row['ClosestFutureDate']

            # Find the line in the text containing the closest future
            closest_date_line = [line for line in deadline_txt.split('\n')
                                 if closest_future_date is not NaT and
                                 closest_future_date.strftime("%d %b %Y") in line]
            closest_date = closest_date_line[0] if closest_date_line else "No upcoming date found"
            p.add_run(f"{closest_date}\n")

        # Amount
        if row['Amount']:
            amount = row['Amount']
            print(f'Summarizing amount: {amount}')
            try:
                response = llm.chat.completions.create(
                    model="gpt-4o",  # Specify the chat model
                    messages=[
                        {"role": "system", "content": "You are a helpful assistant who's good at summarization."},
                        {"role": "user", "content": f"Summarize the following text by extracting award amount in"
                                                f"USD. Is amount upper exists, just use that number is enough."
                                                f"Do not include any notes or explanations:\n\n{amount}"}
                    ],
                    max_tokens=150,  # Set the maximum length for the summary
                    temperature=0.7  # Adjusts randomness in the response. Lower is more deterministic.
                )
                # The response format is different for chat completions
                summary = response.choices[0].message.content.strip()
                print(f'Summarized amount: {summary}')
            except Exception as e:
                print(f"API call failed for amount: {e}")
                summary = amount
            bold_run = p.add_run("Award Amount: ")
            bold_run.bold = True
            p.add_run(f"{summary}\n")

        # Eligibility
        if row['Eligibility']:
            eligibility = format_eligibility(row['Eligibility'], row['Career Stage'])
            print(f'Summarizing eligibility: {eligibility}')

            try:
                messages=[{"role": "system", "content": "You are a helpful, pattern-following assistant that summarizes grant funding opportunities into simple and concise language for UC Irvine faculty.."}]

                # Read the CSV file into a DataFrame
                fse_df = pd.read_csv(FEW_SHOT_EXAMPLES_CSV)

                # Loop over each few-shot example row
                for index, fse_row in fse_df.iterrows():
                    example = format_eligibility(fse_row['Eligibility'], fse_row['Career Stage'])
                    edited = fse_row['Edited']

                    messages.append({"role": "system", "name": "example_user", "content": example})
                    messages.append({"role": "system", "name": "example_assistant", "content": edited})

                # Finally, ask to summarize this grant opportunity
                messages.append({"role": "user", "content": eligibility})

                response = llm.chat.completions.create(
                    model="gpt-4o",  # Specify the chat model
                    messages=messages,
                    max_tokens=200,  # Set the maximum length for the summary
                    temperature=0.7  # Adjusts randomness in the response. Lower is more deterministic.
                )
                summary = response.choices[0].message.content.strip()
                print(f'Summarized eligibility: {summary}')
            except Exception as e:
                print(f"API call failed for Eligibility: {e}")
                summary = eligibility
            # The response format is different for chat completions
            bold_run = p.add_run("Eligibility: ")
            bold_run.bold = True
            p.add_run(f"{summary}\n")

        # Abstract
        if row['Abstract']:
            abstract = row['Abstract']
            print(f'Summarizing abstract: {abstract}')
            try:
                response = llm.chat.completions.create(
                    model="gpt-4o",  # Specify the chat model
                    messages=[
                        {"role": "system", "content": "You are a helpful assistant who's good at summarization."},
                        {"role": "user", "content": f"Summarize the following text in a concise way."
                                                    f"Don’t do the explanation on the foundation itself. Must include "
                                                    f"the specific type of research. Don’t repeat the eligibility nor "
                                                    f"include the budget:\n\n{abstract}"}
                    ],
                    max_tokens=150,  # Set the maximum length for the summary
                    temperature=0.7  # Adjusts randomness in the response. Lower is more deterministic.
                )
                # The response format is different for chat completions
                summary = response.choices[0].message.content.strip()
                print(f'Summarized abstract: {summary}')
            except Exception as e:
                print(f"API call failed for Abstract: {e}")
                summary = abstract
            bold_run = p.add_run("Program Goal: ")
            bold_run.bold = True
            p.add_run(f"{summary}\n")

    # Save the formatted Word document
    # formatted_doc.save(output_file_path)
    return formatted_doc

def save_file(formatted_doc, output_file_path):
    # Save the formatted Word document
    formatted_doc.save(output_file_path)

# Helper functions:
# Add hyperlink
def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a new run for the hyperlink
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    # Optionally, set the formatting (e.g., blue and underlined)
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)

    # Create elements for color and bold
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')  # Blue color
    rPr.append(color)

    bold = OxmlElement('w:b')
    bold.set(qn('w:val'), 'true')
    rPr.append(bold)

    new_run.append(rPr)

    # Add text to the run
    text_element = OxmlElement('w:t')
    text_element.text = text
    new_run.append(text_element)

    # Append the run to the hyperlink element
    hyperlink.append(new_run)

    # Add the hyperlink element to the paragraph
    paragraph._element.append(hyperlink)

    # Create a run object and return it
    return paragraph.add_run()


def extract_closest_future_date(deadline_txt):
    if deadline_txt:
        current_date = datetime.now()
        date_pattern = r"\d{2} \w{3} \d{4}"
        dates = re.findall(date_pattern, deadline_txt)
        future_dates = [datetime.strptime(date, "%d %b %Y") for date in dates if datetime.strptime(date, "%d %b %Y") > current_date]

        if future_dates:
            return min(future_dates, key=lambda x: (x - current_date))
    return None

def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    base_path = getattr(sys, '_MEIPASS', os.path.dirname(os.path.abspath(__file__)))
    return os.path.join(base_path, relative_path)

def get_api_key():
    with open(resource_path(API_KEY_FILE), 'r') as file:
        return file.read().strip()

def unify_line_endings(file_path):
    with open(file_path, 'r', encoding='utf-8', newline=None) as file:
        content = file.read().replace('\r\n', '\n').replace('\r', '\n')

    with open(file_path, 'w', encoding='utf-8', newline='\n') as file:
        file.write(content)
        
def file_process(llm, file_path, head_title):
    unify_line_endings(file_path)
    # 1. read csv file
    data_frame = read_csv_file(file_path)
    # 2. convert csv file to word file and format word file
    if data_frame is not None:
        formatted_doc = format_word_file(llm, data_frame, head_title)
        return formatted_doc
    else:
        raise ValueError("No data found in the file")

def format_eligibility(eligibility, career_stage):
    return eligibility + "\nCareer Stage: " + career_stage

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Process CSV (file or STDIN).")
    parser.add_argument('filename', help="Name of the .csv file to process")
    parser.add_argument('-o', '--output', help="Name of the output file to save the results")
    parser.add_argument('--head', help="The heading title for the document")

    args = parser.parse_args()

    # Default to input filename without .csv if output not provided
    if args.output:
        output_filename = args.output
    else:
        if args.filename.endswith(".csv"):
            output_filename = args.filename[:-4] + ".docx"
        else:
            output_filename = args.filename + ".docx"

    # Default the head (heading) to input filename without .csv and title-cased if not provided
    if args.head:
        head_title = args.head
    else:
        base_name = os.path.basename(args.filename)
        if base_name.endswith(".csv"):
            base_name = base_name[:-4]
        head_title = base_name.replace('_', ' ').title()  # Convert underscores to spaces and title-case it

    llm = AzureOpenAI(
        api_key=get_api_key(),
        api_version="2024-02-01",
        azure_endpoint="https://azureapi.zotgpt.uci.edu/openai/deployments/gpt-4o/chat/completions?api-version=2024-02-01"
    )

    # Process the file and save the results
    formatted_doc = file_process(llm, args.filename, head_title)
    save_file(formatted_doc, output_filename)
